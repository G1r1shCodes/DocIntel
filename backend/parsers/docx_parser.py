"""
DOCX Parser — python-docx with Heading / Table preservation

Preserves heading hierarchies, styled body paragraphs, and embedded
tables. Detects page breaks when present in the document XML.
"""

from __future__ import annotations

import os
import re

from .base import BaseParser, DocumentData, PageData
from .exceptions import ParsingInterruptedError


class DOCXParser(BaseParser):
    """Parse Word documents (``.docx``)."""

    #: Regex that matches the *last rendered page break* marker inserted by
    #: Word.  Used to approximate page boundaries.
    _PAGE_BREAK_RE = re.compile(r'<w:lastRenderedPageBreak\s*/?>', re.IGNORECASE)

    def supported_extensions(self) -> list[str]:
        return [".docx"]  # .doc (OLE) is not supported by python-docx

    def parse(self, file_path: str) -> DocumentData:
        filename = os.path.basename(file_path)
        full_text: list[str] = []
        blocks_data: list[dict] = []
        headings: list[str] = []
        page_count_estimate = 1

        try:
            import docx
        except ImportError:
            raise ParsingInterruptedError(
                "python-docx is not installed; install it with `pip install python-docx`",
                file_path=file_path,
            )

        doc = docx.Document(file_path)

        # --- Paragraphs ---------------------------------------------------
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = (p.style.name if p.style else "").lower()
            if "heading" in style_name:
                headings.append(text)
                blocks_data.append({"text": text, "type": "heading", "style": p.style.name})
            else:
                blocks_data.append({"text": text, "type": "paragraph"})

            full_text.append(text)

        # --- Tables --------------------------------------------------------
        for table_idx, table in enumerate(doc.tables):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_str = "\n".join(rows)
            if table_str.strip():
                full_text.append(f"[TABLE]\n{table_str}\n[/TABLE]")
                blocks_data.append(
                    {"text": table_str, "type": "table", "table_index": table_idx}
                )

        # --- Page-break heuristic (estimate page count) ------------------
        # Look for page-break markers in the document XML
        try:
            body_xml = doc.element.body.xml
            page_breaks = self._PAGE_BREAK_RE.findall(body_xml)
            if page_breaks:
                page_count_estimate = len(page_breaks) + 1
        except Exception:
            pass  # keep estimate of 1

        combined_text = "\n\n".join(full_text)

        page = PageData(
            page_number=1,
            text=combined_text,
            blocks=blocks_data,
            headings=headings,
            is_ocr=False,
        )

        return DocumentData(
            filename=filename,
            file_type="docx",
            text=combined_text,
            metadata={
                "parser": "python-docx",
                "paragraph_count": len(blocks_data),
                "estimated_pages": page_count_estimate,
            },
            pages=[page],
        )
